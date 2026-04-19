"""手动输入密钥"""

import os
import json
import base64
import streamlit as st
from openai import OpenAI
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile

# --------------------------
# 页面基础配置
# --------------------------
# st.set_page_config(
#     page_title="甬高・安眸-施工现场多模态安全隐患识别大模型",
#     page_icon="🛡️",
#     layout="wide",
#     initial_sidebar_state="collapsed"
# )

# 手机端适配
st.html('<meta name="viewport" content="width=device-width, initial-scale=1.0">')

# --------------------------
# 自定义 CSS 样式
# --------------------------
def load_css():
    st.markdown("""
    <style>
        :root {
            --primary-color: #0052D9;
            --primary-dark: #003c9e;
            --danger-color: #E53E3E;
            --warning-color: #FF7D00;
            --success-color: #00B42A;
            --bg-color: #F7F8FA;
            --card-bg: #FFFFFF;
            --text-main: #1D2129;
            --text-secondary: #6E7681;
            --border-color: #E5E6EB;
        }
        .stApp { background-color: var(--bg-color); }
        #MainMenu {visibility: hidden;}
        footer {visibility: hidden;}
        .stDeployButton {display:none;}

        .top-header {
            background: linear-gradient(90deg, #0052D9 0%, #0078F0 100%);
            height: 90px;
            margin: -10px -10px 30px -10px;
            display: flex;
            align-items: center;
            justify-content: center;
            box-shadow: 0 4px 20px rgba(0,82,217,0.15);
        }
        .top-header img { height: 60px; width: auto; }
        .top-title { font-size: 38px; font-weight: 700; color: white; letter-spacing: 1px; }

        .custom-card {
            background: var(--card-bg);
            border-radius: 16px;
            padding: 28px;
            box-shadow: 0 8px 24px rgba(0,0,0,0.06);
            margin-bottom: 24px;
            border: 1px solid rgba(0,0,0,0.05);
            transition: all 0.3s ease;
        }
        .stat-card {
            background: white;
            border-radius: 14px;
            padding: 24px;
            text-align: center;
            box-shadow: 0 4px 16px rgba(0,0,0,0.05);
        }
        .stat-num { font-size: 44px; font-weight: 800; line-height: 1.1; }
        .stat-label { font-size: 15px; color: var(--text-secondary); margin-top: 8px; }
        .tag { display: inline-block; padding: 5px 12px; border-radius: 6px; color: white; font-size: 13px; font-weight: 600; margin-bottom: 10px; }
        .tag-danger { background: linear-gradient(90deg, #E53E3E, #FF4D4F); }
        .tag-warning { background: linear-gradient(90deg, #FF7D00, #FF9500); }
        .upload-area { border: 2px dashed #DCDFE6; border-radius: 12px; padding: 50px 20px; text-align: center; background: white; }
        .hazard-card { border-radius: 12px; padding: 20px; margin-bottom: 18px; background: white; border: 1px solid #F2F3F5; }
    </style>
    """, unsafe_allow_html=True)

# --------------------------
# 核心工具函数
# --------------------------
def image_to_base64(image_source):
    try:
        if isinstance(image_source, str):
            with open(image_source, "rb") as image_file:
                return base64.b64encode(image_file.read()).decode("utf-8")
        else:
            return base64.b64encode(image_source.read()).decode("utf-8")
    except:
        return None

def analyze_image(image_source, api_key):
    image_base64 = image_to_base64(image_source)
    if not image_base64:
        return None

    try:
        client = OpenAI(
            api_key=api_key,
            base_url="https://api.zhizengzeng.com/v1",
        )

        system_prompt = """
        你是一名专业的建筑安全检查员。请分析图片，找出所有安全隐患。如果没有隐患，则不输出隐患。
        请严格以 JSON 格式返回结果，不要包含 Markdown 标记。
        JSON 结构如下：
        {
            "hazards": [
                {
                    "level": "重大" or "一般",
                    "description": "简短描述隐患内容",
                    "confidence": 0.95,
                    "suggestion": "整改建议"
                }
            ]
        }
        """

        completion = client.chat.completions.create(
            model="qwen-vl-plus",
            messages=[
                {"role": "system", "content": system_prompt},
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": "分析这张图片"},
                        {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{image_base64}"}}
                    ]
                }
            ],
            temperature=0.1,
            max_tokens=2000
        )

        res_text = completion.choices[0].message.content
        try:
            if "```json" in res_text:
                res_text = res_text.split("```json")[1].split("```")[0].strip()
            return json.loads(res_text)
        except:
            return {"hazards": [{"level": "一般", "description": res_text, "confidence": 0.9, "suggestion": "请参考描述"}]}
    except Exception as e:
        st.error(f"API 调用失败：{str(e)}")
        return None

def generate_word_report(uploaded_image, result_data):
    doc = Document()
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    def set_chinese_font(run, font_name="微软雅黑", size=12):
        run.font.name = font_name
        run.font.size = Pt(size)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run.font.color.rgb = RGBColor(0,0,0)

    title = doc.add_heading('甬高・安眸-施工现场多模态安全隐患识别大模型检测报告', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        set_chinese_font(run, font_name="微软雅黑", size=18)
        run.bold = True

    img_path = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".jpg", dir=os.getcwd()) as tmp_img:
            tmp_img.write(uploaded_image.getbuffer())
            img_path = tmp_img.name
        doc.add_picture(img_path, width=Inches(5))
    except:
        pass
    finally:
        if img_path and os.path.exists(img_path):
            os.unlink(img_path)

    hazards = result_data.get('hazards', [])
    major = sum(1 for h in hazards if h['level']=='重大')
    minor = sum(1 for h in hazards if h['level']=='一般')

    doc.add_heading('检测统计', level=2)
    p1 = doc.add_paragraph()
    set_chinese_font(p1.add_run(f'重大隐患：{major} 处'))
    p2 = doc.add_paragraph()
    set_chinese_font(p2.add_run(f'一般隐患：{minor} 处'))
    p3 = doc.add_paragraph()
    set_chinese_font(p3.add_run(f'总计：{len(hazards)} 处'))

    doc.add_heading('隐患详情', level=2)
    for i, h in enumerate(hazards, 1):
        t = doc.add_paragraph()
        set_chinese_font(t.add_run(f'{i}. {h["level"]}隐患'), size=13)
        t.runs[0].bold = True
        d = doc.add_paragraph()
        set_chinese_font(d.add_run(f'描述：{h["description"]}'))
        c = doc.add_paragraph()
        set_chinese_font(c.add_run(f'置信度：{int(h.get("confidence",0)*100)}%'))
        s = doc.add_paragraph()
        set_chinese_font(s.add_run(f'建议：{h.get("suggestion","无")}'))

    docx_path = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx', dir=os.getcwd()) as tmp:
            doc.save(tmp.name)
            docx_path = tmp.name
        with open(docx_path, 'rb') as f:
            return f.read()
    finally:
        if docx_path and os.path.exists(docx_path):
            os.unlink(docx_path)

# --------------------------
# 登录页面（必须输密钥才能进）
# --------------------------
def login_page():
    st.markdown("<br><br><br>", unsafe_allow_html=True)
    with st.container(border=True):
        st.title("🔑 请输入 API 密钥")
        api_key = st.text_input("API Key", type="password", label_visibility="collapsed")
        col1, col2, col3 = st.columns([1,2,1])
        with col2:
            if st.button("进入系统", use_container_width=True, type="primary"):
                if api_key and api_key.startswith("sk-"):
                    st.session_state["api_key"] = api_key
                    st.rerun()
                else:
                    st.error("请输入有效的 API 密钥")

# --------------------------
# 主界面
# --------------------------
def main_page():
    load_css()
    api_key = st.session_state.get("api_key", "")

    if 'result_data' not in st.session_state:
        st.session_state['result_data'] = None
    if 'uploaded_file' not in st.session_state:
        st.session_state['uploaded_file'] = None

    def get_base64_of_bin_file(bin_file):
        with open(bin_file, 'rb') as f:
            return base64.b64encode(f.read()).decode()

    try:
        logo_base64 = get_base64_of_bin_file("logo1.png")
    except:
        logo_base64 = ""

    st.markdown(f"""
    <div class="top-header">
        <img src="data:image/png;base64,{logo_base64}">
        <span class="top-title"> 甬高・安眸 — 施工现场多模态安全隐患识别大模型</span>
    </div>
    """, unsafe_allow_html=True)

    col1, col2 = st.columns([1.2,1], gap="large")

    with col1:
        st.markdown('<div class="custom-card">', unsafe_allow_html=True)
        st.markdown("#### 📷 上传识别")
        uploaded_file = st.file_uploader(" ", type=["jpg","jpeg","png"], label_visibility="collapsed")
        if uploaded_file:
            st.session_state['uploaded_file'] = uploaded_file
            st.image(uploaded_file, caption="上传图片", use_column_width=True)
        else:
            st.markdown('<div class="upload-area"><div style="font-size:48px">📷</div><div>上传施工现场图片</div></div>', unsafe_allow_html=True)
        st.markdown('</div>', unsafe_allow_html=True)

        if st.button("开始识别", use_container_width=True, type="primary"):
            if uploaded_file:
                with st.spinner("正在分析..."):
                    res = analyze_image(uploaded_file, api_key)
                    st.session_state['result_data'] = res

    with col2:
        st.markdown('<div class="custom-card">', unsafe_allow_html=True)
        st.markdown("### 📊 检测概览")
        s1, s2 = st.columns(2)
        major = 0
        minor = 0
        if st.session_state['result_data']:
            hh = st.session_state['result_data'].get('hazards',[])
            major = sum(1 for h in hh if h['level']=='重大')
            minor = sum(1 for h in hh if h['level']=='一般')
        with s1:
            st.markdown(f'<div class="stat-card"><div class="stat-num" style="color:red">{major}</div><div class="stat-label">重大隐患</div></div>', unsafe_allow_html=True)
        with s2:
            st.markdown(f'<div class="stat-card"><div class="stat-num" style="color:orange">{minor}</div><div class="stat-label">一般隐患</div></div>', unsafe_allow_html=True)
        st.markdown('</div>', unsafe_allow_html=True)

        st.markdown("### 📋 隐患详情")
        if st.session_state['result_data']:
            hz = st.session_state['result_data'].get('hazards',[])
            if not hz:
                st.success("✅ 未发现安全隐患")
            else:
                for h in hz:
                    tag = "tag-danger" if h['level']=="重大" else "tag-warning"
                    st.markdown(f'''
                    <div class="hazard-card">
                        <span class="tag {tag}">{h['level']}隐患</span>
                        <div style="font-weight:bold; margin-bottom:6px;">{h['description']}</div>
                        <div style="font-size:13px; color:#666">
                        置信度：{int(h.get('confidence',0)*100)}%<br>
                        整改建议：{h.get('suggestion','无')}
                        </div>
                    </div>
                    ''', unsafe_allow_html=True)
        else:
            st.info("请上传图片并开始识别")

        if st.session_state['result_data'] and st.session_state['uploaded_file']:
            b = generate_word_report(st.session_state['uploaded_file'], st.session_state['result_data'])
            st.download_button("📄 导出Word报告", b, f"安全检测报告.docx", use_container_width=True)

# --------------------------
# 路由控制
# --------------------------
if "api_key" not in st.session_state or not st.session_state["api_key"]:
    login_page()
else:
    main_page()